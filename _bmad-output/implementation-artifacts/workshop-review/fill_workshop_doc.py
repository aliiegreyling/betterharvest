from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOC_PATH = Path(
    "/Users/alexgreyling/Documents/betterharvest/"
    "_bmad-output/implementation-artifacts/workshop-review/"
    "McConference2026 WorkshopReview.docx"
)
OUT_PATH = DOC_PATH.with_name("McConference2026 WorkshopReview - Code Forgers.docx")

TEAM_MEMBERS = [
    ("Alex Greyling", "Team member"),
    ("Brandon van Heerden", "Team member"),
    ("Connor Cress", "Team Lead / Point of Contact"),
    ("Connor Foley", "Team member"),
    ("Daniel van Tonder", "Team member"),
    ("Divan Meyer", "Team member"),
    ("Gerhard Kloppers", "Team member"),
    ("Keelan Matthews", "Team member"),
    ("Kyle Marshall", "Team member"),
    ("Michael Rademeyer", "Team member"),
    ("Michal van der Merwe", "Team member"),
    ("Muller Dannhauser", "Team member"),
    ("Nicholas Naicker", "Team member"),
    ("Reinhart du Plessis", "Team member"),
    ("Tiaan Pouwels", "Team member"),
    ("Tian Bornman", "Team member"),
]


def copy_paragraph_format(source: Paragraph, target: Paragraph) -> None:
    target.style = source.style
    target.alignment = source.alignment
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.space_before = source.paragraph_format.space_before
    target.paragraph_format.space_after = source.paragraph_format.space_after
    target.paragraph_format.line_spacing = source.paragraph_format.line_spacing


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    copy_paragraph_format(paragraph, new_para)
    if paragraph._p.pPr is not None:
        new_para._p.insert(0, deepcopy(paragraph._p.pPr))
    new_para.add_run(text)
    return new_para


doc = Document(DOC_PATH)

set_paragraph_text(doc.paragraphs[5], "Team Name: Code Forgers")
set_paragraph_text(doc.paragraphs[6], "Team Members (Names & Roles):")

anchor = doc.paragraphs[6]
for name, role in reversed(TEAM_MEMBERS):
    insert_paragraph_after(anchor, f"- {name} - {role}")

doc = Document(DOC_PATH)
set_paragraph_text(doc.paragraphs[5], "Team Name: Code Forgers")
set_paragraph_text(doc.paragraphs[6], "Team Members (Names & Roles):")

anchor = doc.paragraphs[6]
for name, role in TEAM_MEMBERS:
    anchor = insert_paragraph_after(anchor, f"- {name} - {role}")

doc = Document(DOC_PATH)
set_paragraph_text(doc.paragraphs[5], "Team Name: Code Forgers")
set_paragraph_text(doc.paragraphs[6], "Team Members (Names & Roles):")

anchor = doc.paragraphs[6]
for name, role in TEAM_MEMBERS:
    anchor = insert_paragraph_after(anchor, f"- {name} - {role}")

set_paragraph_text(doc.paragraphs[8 + len(TEAM_MEMBERS)], "Team Lead / Point of Contact: Connor Cress")
set_paragraph_text(doc.paragraphs[9 + len(TEAM_MEMBERS)], "Date: May 21, 2026")
set_paragraph_text(
    doc.paragraphs[95 + len(TEAM_MEMBERS)],
    "Submitted by: Connor Cress     Time: __________",
)

doc.save(OUT_PATH)
print(OUT_PATH)
